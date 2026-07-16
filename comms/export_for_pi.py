#!/usr/bin/env python3
"""
export_for_pi.py — generate the PI-facing pack from the repo's own truth.
NEVER edit the outputs by hand: edit research_log.md / figure_register.md /
make commits, then re-run this. Outputs land in comms/PI_pack/.

  Progress_Note.docx      <- latest N entries of research_log.md
  Revision_History.docx   <- git log (date, hash, message)
  Figure_Register.xlsx    <- comms/figure_register.md table

Run from the repo root:  python comms/export_for_pi.py
"""
import os, re, subprocess, datetime

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT  = os.path.join(ROOT, "comms", "PI_pack")
os.makedirs(OUT, exist_ok=True)
N_ENTRIES = 6   # how many latest log entries go into the progress note

from docx import Document
from docx.shared import Pt, RGBColor, Inches
import openpyxl
from openpyxl.styles import Font, PatternFill

NAVY = RGBColor(0x1F, 0x38, 0x64)

def styled_doc(title, subtitle):
    d = Document()
    for sec in d.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.8)
    h = d.add_heading(title, 0)
    for r in h.runs: r.font.color.rgb = NAVY
    p = d.add_paragraph(subtitle)
    p.runs[0].font.size = Pt(10); p.runs[0].font.italic = True
    return d

# ---------------- 1. Progress note from research_log.md -------------------
log = open(os.path.join(ROOT, "research_log.md"), encoding="utf-8").read()
entries = re.split(r"\n(?=## \d{4}-\d{2}-\d{2})", log)
entries = [e for e in entries if e.startswith("## 2")][-N_ENTRIES:][::-1]  # newest first

d = styled_doc("Progress Note — Shovel Load-Torque Project",
    f"Project IIT/SRIC/R/AEH/2026/104 · generated {datetime.date.today()} from research_log.md · JRF: Suhail Majeed Sheikh")
for e in entries:
    lines = e.strip().splitlines()
    head = lines[0].lstrip("# ").strip()
    h = d.add_heading(head, level=1)
    for r in h.runs: r.font.color.rgb = NAVY; r.font.size = Pt(13)
    for ln in lines[1:]:
        ln = ln.rstrip()
        if not ln: continue
        if ln.startswith("**") and ln.rstrip(":*").endswith(("done", "result", "why", "step", "questions")) or re.match(r"\*\*[^*]+\*\*", ln):
            m = re.match(r"\*\*([^*]+)\*\*:?\s*(.*)", ln)
            if m:
                para = d.add_paragraph()
                run = para.add_run(m.group(1) + ": "); run.bold = True
                para.add_run(m.group(2))
                continue
        d.add_paragraph(re.sub(r"\*\*", "", ln))
d.save(os.path.join(OUT, "Progress_Note.docx"))

# ---------------- 2. Revision history from git -----------------------------
gitlog = subprocess.run(["git", "-C", ROOT, "log", "--date=short",
    "--pretty=format:%ad|%h|%s"], capture_output=True, text=True).stdout
d = styled_doc("Revision History",
    f"Generated {datetime.date.today()} from the git repository (github.com/sonu037/shovel_dynamics)")
t = d.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
for c, txt in zip(t.rows[0].cells, ("Date", "Version", "Change")):
    c.text = txt
    c.paragraphs[0].runs[0].font.bold = True
for line in gitlog.splitlines():
    date, h, msg = line.split("|", 2)
    row = t.add_row().cells
    row[0].text, row[1].text, row[2].text = date, h, msg
d.save(os.path.join(OUT, "Revision_History.docx"))

# ---------------- 3. Figure register xlsx from md --------------------------
md = open(os.path.join(ROOT, "comms", "figure_register.md"), encoding="utf-8").read()
rows = [r for r in md.splitlines() if r.startswith("|")]
table = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
table = [r for r in table if not set("".join(r)) <= set("-: ")]  # drop separator

wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Figure Register"
navy = PatternFill("solid", fgColor="1F3864")
for j, cell in enumerate(table[0], 1):
    c = ws.cell(row=1, column=j, value=cell)
    c.font = Font(bold=True, color="FFFFFF"); c.fill = navy
for i, r in enumerate(table[1:], 2):
    for j, cell in enumerate(r, 1):
        ws.cell(row=i, column=j, value=cell)
widths = [9, 34, 10, 8, 11, 10, 42, 46]
for j, w in enumerate(widths, 1):
    ws.column_dimensions[openpyxl.utils.get_column_letter(j)].width = w
wb.save(os.path.join(OUT, "Figure_Register.xlsx"))

print("PI pack written to", OUT)
